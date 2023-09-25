from flask import Flask, request, jsonify
from plantuml import PlantUML
from docx import Document
from docx.shared import Inches
import openai
import json
import random
import string
import os
from azure.storage.blob import BlobServiceClient, BlobClient, ContainerClient

openai.api_key = ''
model_name = 'gpt-3.5-turbo'
template_dict = {
   '1' : """
>>><input_text><<<

======
Intsuction prompt
Read >>>CONTENT<<< and generate JSON in the below format
{
"DesignDoc": {
  "Title":,
  "Intent":,
  "Solution",
  "Pros",
  "Cons",}.
"Diagram": this provides diagram of the content in plantuml output format
}
======
Please ensure that the output is in a perfect JSON format which can be serialized and deserialized.
---
Example

User:I have a service A, which calls service B for sku info, and service C for OS info. Based on information from B and C, service A calls a method called pushConfig which takes sku and os as input. Based on sku and os, the pushConfig method sends a command to a device.
Assistant: Here is the data you requested:
{
"DesignDoc":{
  "Title":,
  "Intent",
  "Solution",
  "Pros",
  "Cons",
  "Detailed System Design"},
"Diagram": \"\"\"
@startuml
actor User

User -> ServiceA : Request SKU Info
ServiceA -> ServiceB : Request SKU Info
ServiceB -> ServiceA : Return SKU Info
ServiceA -> ServiceC : Request OS Info
ServiceC -> ServiceA : Return OS Info
ServiceA -> Device : Send Command

@enduml
\"\"\"
}
""",
'2' : """ 
>>><input_text><<<

======
Intsuction prompt
Read >>>CONTENT<<< and generate JSON in the below format(Use this template if you want to generate a design document for a new service, ensure that the result has all the fields mentioned below)
{
"DesignDoc": {
  "Title": Title of the design document,
  "Introduction" : Provide an overview of the entire document,
  "System Overview" : Provide a general description and functionality of the software system. ,
  "Design Considerations": Describe the issues that need to be addressed before creating a design solution,
  "Assumptions and Dependencies" : Describe any assumptions that may be wrong or any dependencies on other things,
  "System Architecture" : This section should provide a high-level overview of how the functionality and responsibilities of the system were partitioned and then assigned to subsystems or components.,
  "Policies and Tactics": Describe any design policies and/or tactics that do not have sweeping architectural implications (meaning they would not significantly affect the overall organization of the system and its high-level structures),
  "Pseudo Code": This section should provide a high-level implementation in psuedo-code of how the functionality and responsibilities of the system were partitioned and then assigned to subsystems or components.,
  "Detailed System Design": Most components described in the System Architecture section will require a more detailed discussion. Other lower-level components and subcomponents may need to be described as well.,
  "Glossary": An ordered list of defined terms and concepts used throughout the document.},
"Diagram": this provides diagram of the content in plantuml output format
}
======
Please ensure that the output is in a perfect JSON format which can be serialized and deserialized. Please ensure that system overview section and detailed system design section are atleast 100 words long
---
Example

User:I have a service A, which calls service B for sku info, and service C for OS info. Based on information from B and C, service A calls a method called pushConfig which takes sku and os as input. Based on sku and os, the pushConfig method sends a command to a device.
Assistant: Here is the data you requested:
{
"DesignDoc": {
  "Title":,
  "Introduction",
  "System Overview",
  "Design Considerations",
  "Assumptions and Dependencies",
  "System Architecture",
  "Policies and Tactics",
  "Pseudo Code",
  "Detailed System Design",
  "Glossary"},
"Diagram": \"\"\"
@startuml
actor User

User -> ServiceA : Request SKU Info
ServiceA -> ServiceB : Request SKU Info
ServiceB -> ServiceA : Return SKU Info
ServiceA -> ServiceC : Request OS Info
ServiceC -> ServiceA : Return OS Info
ServiceA -> Device : Send Command

@enduml
\"\"\"
}
"""
}
uml_text_file_path = "C:\Hack23\ChatDesignDoc\plantUml.txt"
uml_img_file_path = "C:\Hack23\ChatDesignDoc\plantUml.png"
doc_path = 'C:\Hack23\ChatDesignDoc\chatgptdoc.docx'
blob_service_client = BlobServiceClient.from_connection_string("BlobEndpoint=https://testingaci.blob.core.windows.net/;QueueEndpoint=https://testingaci.queue.core.windows.net/;FileEndpoint=https://testingaci.file.core.windows.net/;TableEndpoint=https://testingaci.table.core.windows.net/;SharedAccessSignature=sv=2022-11-02&ss=bfqt&srt=sco&sp=rwdlacupiytfx&se=2024-01-25T02:57:58Z&st=2023-09-13T18:57:58Z&spr=https&sig=66XZYbaPFUROK4CI%2BAWP1%2BRghH%2BsfQKawjlnLoiZS98%3D")
container_name = 'hackchat'

# Functions

def replace_input_text(input_string, template):
    return template.replace("<input_text>", input_string)

def get_completion(template_enum, input_text):
    if template_enum is None:
        template_enum = "1"
    template = template_dict[template_enum]
    input_text_in_template = replace_input_text(input_text, template)
    print(input_text_in_template)
    chat_completion = openai.ChatCompletion.create(model="gpt-3.5-turbo", messages=[{"role": "user", "content": input_text_in_template}])
    return chat_completion.choices[0].message.content

def Convert_string_to_json(st):
  st = st.replace('"""','"')
  json_output = json.loads(st, strict=False)
  return json_output

def write_file(text, path):
  with open(path, 'w') as f:
    f.write(text)

def Create_plantUml_img(source, destination):
  server = PlantUML(url='http://www.plantuml.com/plantuml/img/',
                            basic_auth={},
                            form_auth={}, http_opts={}, request_opts={})
  server.processes_file(source, destination)

def create_document_from_json(json_obj, main_heading, img_path, doc_path):
  document = Document()
  document.add_heading(main_heading, 0)

  for key in json_obj.keys():
    if(key != 'Title'):
      document.add_heading(key, level=1)
      
      if(isinstance(json_obj[key], str)):
        document.add_paragraph(json_obj[key])

      if(isinstance(json_obj[key], list)):
        for list_item in json_obj[key]:
          document.add_paragraph(list_item, style='List Bullet')

  document.add_picture(img_path, width=Inches(3.5))

  document.add_page_break()

  document.save(doc_path)


def generate_random_string(size):
    chars = string.ascii_lowercase
    return ''.join(random.choice(chars) for _ in range(size))

def add_file_to_blob(file_path, blob_name, container_name):
    # Create a blob client using the local file name as the name for the blob
    container_client = blob_service_client.get_container_client(container_name)
    blob_client = container_client.get_blob_client(blob_name)

    # Upload the created file
    with open(file_path, "rb") as data:
        blob_client.upload_blob(data)
    return blob_client.url

def retry_logic(template_input, text_input, retry_success = 0):
    if(retry_success == 2):
      return {'document': "Retry failed", 'diagram' : "Retry failed", 'input_text': text_input, 'output_gpt': "Retry failed"}
    output_gpt = get_completion(template_input, text_input)
    try:
      json_output_gpt = Convert_string_to_json(output_gpt)
      write_file(json_output_gpt['Diagram'], uml_text_file_path)
      Create_plantUml_img(uml_text_file_path, uml_img_file_path)
      create_document_from_json(json_output_gpt['DesignDoc'], json_output_gpt['DesignDoc']['Title'], uml_img_file_path, doc_path)
      output_doc_blob_link = add_file_to_blob(doc_path, generate_random_string(10)+".docx", container_name)
      output_img_blob_link = add_file_to_blob(uml_img_file_path, generate_random_string(10)+".png", container_name)
      return {'document': output_doc_blob_link, 'diagram' : output_img_blob_link, 'input_text': text_input, 'output_gpt': json_output_gpt}
    except:
      return retry_logic(template_input, text_input, retry_success+1)
      # return {'document': output_doc_blob_link, 'diagram' : output_img_blob_link, 'input_text': text_input, 'output_gpt': json_output_gpt}
   
app = Flask(__name__)

@app.route('/gpt', methods=['GET'])
def gpt():
    text_input = request.args.get('text_input')
    template_input = request.args.get('template', None)
    response = retry_logic(template_input, text_input)
    return jsonify(response)

if __name__ == '__main__':
    app.run(host='0.0.0.0', debug=True)
